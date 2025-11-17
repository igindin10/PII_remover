import hashlib
import re
import string
import sys
from pathlib import Path

import pandas as pd
import phonenumbers
from bs4 import BeautifulSoup, Comment
from bs4.element import AttributeValueList
from docx import Document
from presidio_analyzer import AnalyzerEngine, RecognizerRegistry, PatternRecognizer, Pattern
from presidio_analyzer.nlp_engine import SpacyNlpEngine
from presidio_anonymizer import AnonymizerEngine, OperatorConfig
from wordfreq import zipf_frequency
from pdfminer.high_level import extract_text


class InternationalPhoneNumberRecognizer(PatternRecognizer):
    def __init__(self):
        # A permissive regex pattern to catch various international phone formats
        # Matches +<digits> with optional spaces, dashes, or parentheses
        pattern = Pattern(
            name="international_phone_number",
            regex=r"\+?\d{1,3}[\s.-]?\d{1,3}[\s.-]?\d{3}[\s.-]?\d{4}(?!\d)",
            score=0.7
        )
        super().__init__(
            supported_entity="PHONE_NUMBER",
            patterns=[pattern]
        )

    def validate_result(self, pattern_text: str) -> bool:
        """
        Overrides the default validator to use the phonenumbers library.
        Returns True only if phonenumbers.is_valid_number() passes.
        """
        try:
            # Try to parse the number with no region bias (None)
            num = phonenumbers.parse(pattern_text, None)
            return phonenumbers.is_valid_number(num)
        except Exception:
            return False



UNICODE_SPACE_MAP = {
    "\u00A0": " ",  # NBSP
    "\u202F": " ",  # NARROW NBSP
    "\u2007": " ",  # FIGURE SPACE
}
UNICODE_DASH_MAP = {
    "\u2010": "-",  # HYPHEN
    "\u2011": "-",  # NON-BREAKING HYPHEN
    "\u2012": "-",  # FIGURE DASH
    "\u2013": "-",  # EN DASH
    "\u2014": "-",  # EM DASH
    "\u2212": "-",  # MINUS SIGN
}

def _normalize_unicode_separators(s: str) -> str:
    # fast path: only do replaces if needed
    for u, r in UNICODE_SPACE_MAP.items():
        if u in s:
            s = s.replace(u, r)
    for u, r in UNICODE_DASH_MAP.items():
        if u in s:
            s = s.replace(u, r)
    return s


# === Presidio / redaction config ===
REDACT_ENTITIES = [
    "EMAIL_ADDRESS",
    "PHONE_NUMBER",
    "PERSON",
    "IP_ADDRESS",
    "CREDIT_CARD",
    "IBAN",
    "US_SSN",
    "TOKEN",
    # Can be added more as needed
]

ANALYZER_THRESHOLD = 0.4

#Anonymizer operators
ANONYMIZER_OPERATORS = {
    "DEFAULT": OperatorConfig("replace", {"new_value": "[REDACTED]"}),
    "EMAIL_ADDRESS": OperatorConfig("replace", {"new_value": "[EMAIL ADDRESS REDACTED]"}),
    "PHONE_NUMBER": OperatorConfig("replace", {"new_value": "[PHONE NUMBER REDACTED]"}),
    "PERSON": OperatorConfig("replace", {"new_value": "[PERSON REDACTED]"}),
    "IP_ADDRESS": OperatorConfig("replace", {"new_value": "[IP ADDRESS REDACTED]"}),
    "CREDIT_CARD": OperatorConfig("replace", {"new_value": "[CREDIT CARD REDACTED]"}),
    "IBAN": OperatorConfig("replace", {"new_value": "[IBAN REDACTED]"}),
    "US_SSN": OperatorConfig("replace", {"new_value": "[US SSN REDACTED]"}),
    "TOKEN": OperatorConfig("replace", {"new_value": "[TOKEN REDACTED]"}),
}

# === Configuration ===
INPUT_PATH = sys.argv[1] if len(sys.argv) > 1 else "."
OUTPUT_DIR = "redacted"
LOG_DIR = "logs"
SUPPORTED_EXTENSIONS = [".txt", ".csv", ".html", ".htm", ".docx", "pdf"]

RE_EMAIL = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
RE_TOKEN = re.compile(r"[a-zA-Z0-9]{10,}[.!][a-zA-Z0-9._-]{5,}[.!][a-zA-Z0-9._-]{5,}")
RE_SSN = re.compile(r"\b\d{3}-\d{2}-\d{4}\b")
RE_IBAN = re.compile(r"\b[A-Z]{2}\d{2}(?: ?[A-Z0-9]){11,30}\b")
CONTEXT_POS = re.compile(r"(?i)\b(to|from|attn|cc|bcc|dear|regards|sincerely|signed|contact)\b")
CONTEXT_NEG = re.compile(r"(?i)\b(email|reason|status|owner|loading|workbook|user|name|value)\b")
_NAME_SEP = r"(?: |\u00A0|\u202F|\u2007)"
_FIRST_LAST = re.compile(rf"^[A-Z][a-z]+{_NAME_SEP}[A-Z][a-z]+$")
RE_FIRST_LAST = re.compile(rf"\b[A-Z][a-z]+{_NAME_SEP}[A-Z][a-z]+\b")
RE_E164_LAX = re.compile(r"\+\d[\d\-\s()]{6,}\d")

# Do not redact these suspicious "person" labels
FALSE_POSITIVE_PERSONS = {s.lower() for s in {
    "email", "reason", "loading", "workbook", "status", "owner", "assigned",
    "priority", "default", "total", "value", "name", "user", "users"
}}


def build_presidio():
    models = [
        {"lang_code": "en", "model_name": "en_core_web_lg"},
        #{"lang_code": "he", "model_name": "he_core_news_lg"},  # for future usage
    ]
    nlp_engine = SpacyNlpEngine(models=models)
    try:
        nlp_engine.load()
    except Exception as ex:
        raise RuntimeError(
            "Failed to load spaCy model(s). Ensure `python -m spacy download en_core_web_lg` "
            "was run inside this virtualenv.\nOriginal error: " + str(ex)
        )

    registry = RecognizerRegistry()
    registry.load_predefined_recognizers(nlp_engine=nlp_engine)

    token_recognizer = PatternRecognizer(
        name="TOKEN_RECOGNIZER",
        supported_entity="TOKEN",
        patterns=[Pattern(name="token", regex=RE_TOKEN.pattern, score=0.8)],
    )

    registry.add_recognizer(InternationalPhoneNumberRecognizer())
    registry.add_recognizer(token_recognizer)

    analyzer = AnalyzerEngine(nlp_engine=nlp_engine, registry=registry)
    anonymizer = AnonymizerEngine()
    return analyzer, anonymizer

ANALYZER, ANONYMIZER = build_presidio()

def save_redacted_text(text, output_path):
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)

def is_suspected_name(t: str) -> bool:
    t = t.strip().strip(string.punctuation)
    if not t:
        return False
    if t.lower() in FALSE_POSITIVE_PERSONS:
        return False
    #avoid obvious non-names
    if len(t.split()) > 3 or t.isupper() or t.isnumeric():
        return False
    if _FIRST_LAST.match(t):
        return True
    #common words filter (zipf>=4.5 are common)
    words = re.findall(r"[A-Za-z]+", t)
    if any(zipf_frequency(w.lower(), "en") >= 4.5 for w in words):
        return False
    return True

def should_redact_person(span_text: str, left_ctx: str, right_ctx: str) -> bool:
    if CONTEXT_NEG.search(left_ctx) or CONTEXT_NEG.search(right_ctx):
        return False
    if CONTEXT_POS.search(left_ctx) or CONTEXT_POS.search(right_ctx):
        return True
    return is_suspected_name(span_text)

def redact_phones_with_validator(text: str, log=None) -> str:
    """
    Redact phone numbers in local (IL) and international (+<cc>) formats.
    Tries region='IL' first for local numbers, then falls back to None for global.
    """
    # Normalize Unicode dashes → simple hyphen
    text_norm = text.replace("–", "-").replace("—", "-")

    out = []

    matches = []

    # Pass 1: local Israeli numbers (region='IL')
    for m in phonenumbers.PhoneNumberMatcher(
        text_norm, "IL", leniency=phonenumbers.Leniency.POSSIBLE
    ):
        if phonenumbers.is_valid_number(m.number):
            matches.append((m.start, m.end))

    # Pass 2: international +country code numbers (region=None)
    for m in phonenumbers.PhoneNumberMatcher(
        text_norm, None, leniency=phonenumbers.Leniency.POSSIBLE
    ):
        if phonenumbers.is_valid_number(m.number):
            matches.append((m.start, m.end))

    # Merge overlaps (rare but safe)
    matches = sorted(set(matches))
    last_end = 0
    for start, end in matches:
        if start < last_end:  # overlapping match, skip
            continue
        out.append(text_norm[last_end:start])
        span = text_norm[start:end]
        if log is not None:
            log.append(("PHONE_NUMBER", span))
        out.append("[PHONE REDACTED]")
        last_end = end
    out.append(text_norm[last_end:])

    return "".join(out)



def _fallback_redact_e164(text: str, log: list) -> str:
    out = []
    i = 0
    for m in RE_E164_LAX.finditer(text):
        start, end = m.start(), m.end()
        out.append(text[i:start])
        cand = m.group(0)
        try:
            num = phonenumbers.parse(cand, None)  # None: infer from +<cc>
            if phonenumbers.is_valid_number(num):
                log.append(("PHONE_NUMBER", cand))
                out.append("[PHONE REDACTED]")
            else:
                out.append(cand)
        except Exception:
            out.append(cand)
        i = end
    out.append(text[i:])
    return "".join(out)


def redact_text(text: str, log: list) -> str:
    def log_and_replace(pattern, replacement, label_name):
        matches = list(pattern.finditer(text))
        for match in matches:
            log.append((label_name, match.group()))
        return pattern.sub(replacement, text)

    text = _normalize_unicode_separators(text)

    text = log_and_replace(RE_EMAIL, "[EMAIL REDACTED]", "EMAIL_ADDRESS")
    text = log_and_replace(RE_TOKEN, "[TOKEN REDACTED]", "TOKEN")
    text = log_and_replace(RE_SSN, "[US SSN REDACTED]", "US_SSN")
    text = log_and_replace(RE_IBAN, "[IBAN REDACTED]", "IBAN")
    text = redact_phones_with_validator(text, log=log)
    text = _fallback_redact_e164(text, log)

    #run Presidio analyzer
    results = ANALYZER.analyze(
        text=text,
        entities=REDACT_ENTITIES,
        language="en",
        score_threshold=ANALYZER_THRESHOLD,
    )

    #anonymize with deterministic replacements
    anonymized = ANONYMIZER.anonymize(
        text=text,
        # analyzer_results=filtered,
        analyzer_results=results,
        operators=ANONYMIZER_OPERATORS
    ).text

    for r in results:
        span = text[r.start:r.end]
        log.append((r.entity_type, span))

    return anonymized

def process_txt(file_path_txt: Path, out_path: Path, log_path: Path):
    log = []
    with open(file_path_txt, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    redacted = redact_text(text, log)
    save_redacted_text(redacted, out_path)
    write_log(log, log_path)

def process_csv(file_path_csv: Path, out_path: Path, log_path: Path):
    def _redact_cell(x, log_list):
        if pd.isna(x):
            return x
        s = str(x).strip()
        return redact_text(s, log_list)

    log = []
    df = pd.read_csv(file_path_csv, dtype=str, encoding="utf-8", on_bad_lines="skip")
    for col in df.columns:
        df[col] = df[col].apply(lambda v: _redact_cell(v, log))
    df.to_csv(out_path, index=False)
    write_log(log, log_path)


def process_html(file_path_html: Path, out_path: Path, log_path: Path):
    log = []
    with open(file_path_html, "r", encoding="utf-8", errors="ignore") as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    #redact a visible text
    for node in soup.find_all(string=True):
        if node.parent.name in {"script", "style"} or isinstance(node, Comment):
            continue
        redacted = redact_text(str(node), log)
        if redacted != str(node):
            node.replace_with(soup.new_string(redacted))

    #redact inside element attributes
    pii_attrs = {"href", "title", "alt", "content", "data-email", "data-user"}
    for tag in soup.find_all(True):  # all element tags
        for attr, val in list(tag.attrs.items()):
            #list-valued attributes
            if isinstance(val, list):
                new_list, changed = [], False
                for v in val:
                    if isinstance(v, str):
                        red = redact_text(v, log)
                        if red != v:
                            changed = True
                        new_list.append(red)
                    else:
                        new_list.append(v)
                if changed:
                    if attr in {"class", "rel"}:
                        tag[attr] = AttributeValueList(new_list)
                    else:
                        tag[attr] = " ".join(new_list)
                continue
            #string attributes (filter to likely-PII attributes or any data-*)
            if isinstance(val, str) and (attr in pii_attrs or attr.startswith("data-")):
                red = redact_text(val, log)
                if red != val:
                    tag[attr] = red

    with open(out_path, "w", encoding="utf-8") as f:
        f.write(str(soup))
    write_log(log, log_path)

def process_docx(file_path_doc: Path, out_path: Path, log_path: Path):
    log = []
    doc = Document(str(file_path_doc))

    #redact regular paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = redact_text(run.text, log)

    #redact text inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.text:
                            r.text = redact_text(r.text, log)

    #redact headers and footers
    for section in doc.sections:
        #header
        for p in section.header.paragraphs:
            for r in p.runs:
                if r.text:
                    r.text = redact_text(r.text, log)
        #footer
        for p in section.footer.paragraphs:
            for r in p.runs:
                if r.text:
                    r.text = redact_text(r.text, log)

    doc.save(str(out_path))
    write_log(log, log_path)

def process_pdf_to_txt(file_path_pdf: Path, out_path: Path, log_path: Path):
    log = []
    text = extract_text(file_path_pdf)
    redacted = redact_text(text, log)
    out_path = out_path.with_suffix(".txt")
    save_redacted_text(redacted, out_path)
    write_log(log, log_path)

def write_log(log_entries, log_path):
    def _mask_for_log(x: str) -> str:
        h = hashlib.sha256(x.encode("utf-8")).hexdigest()[:8]
        #small preview for debugging only; must be removed if a policy forbids
        return f"<{x[:3]}…#{h}>"

    with open(log_path, "w", encoding="utf-8") as f:
        for label, value in log_entries:
            f.write(f"{label} => {_mask_for_log(value)}\n")

Path(OUTPUT_DIR).mkdir(exist_ok=True)
Path(LOG_DIR).mkdir(exist_ok=True)

input_path = Path(INPUT_PATH)
files = []
if input_path.is_dir():
    for ext in SUPPORTED_EXTENSIONS:
        files.extend(input_path.rglob(f"*{ext}"))
elif input_path.suffix.lower() in SUPPORTED_EXTENSIONS:
    files = [input_path]
else:
    print("Unsupported file or directory.")
    sys.exit(1)

for file_path in files:
    file_path = Path(file_path)
    name, ext = file_path.stem, file_path.suffix.lower()
    out_file = Path(OUTPUT_DIR) / f"{name}_redacted{ext}"
    log_file = Path(LOG_DIR) / f"{name}_log.txt"
    try:
        if ext == ".txt":
            process_txt(file_path, out_file, log_file)
        elif ext == ".csv":
            process_csv(file_path, out_file, log_file)
        elif ext in [".html", ".htm"]:
            process_html(file_path, out_file, log_file)
        elif ext == ".docx":
            process_docx(file_path, out_file, log_file)
        elif ext == ".pdf":
            process_pdf_to_txt(file_path, out_file, log_file)
        print(f"Redacted: {file_path.name}")
    except Exception as e:
        print(f"Error processing {file_path.name}: {e}")
